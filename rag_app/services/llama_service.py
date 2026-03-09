import os
from functools import lru_cache
from django.conf import settings
from llama_index.core import VectorStoreIndex, StorageContext, SimpleDirectoryReader
from llama_index.vector_stores.postgres import PGVectorStore

# 尝试导入 DashScope Rerank
try:
    from llama_index.postprocessors.dashscope_rerank import DashScopeRerank
    HAS_RERANK = True
except ImportError:
    HAS_RERANK = False
    DashScopeRerank = None

# 尝试直接使用 DashScope API 进行 rerank
try:
    import dashscope
    from dashscope import TextReRank
    HAS_DASHSCOPE_RERANK = True
except ImportError:
    HAS_DASHSCOPE_RERANK = False
    TextReRank = None


# 缓存 vector index 和 query engine 实例
_vector_index_cache = {}
_query_engine_cache = {}


def get_query_engine(tenant_id: str):
    """
    获取缓存的 Query Engine，避免每次查询都创建新实例
    """
    cache_key = f"qe_{tenant_id}"
    if cache_key in _query_engine_cache:
        return _query_engine_cache[cache_key]

    index = get_vector_index()

    # Data Isolation: Only retrieve context matching the tenant_id
    from llama_index.core.vector_stores import MetadataFilters, ExactMatchFilter
    filters = MetadataFilters(
        filters=[ExactMatchFilter(key="tenant_id", value=tenant_id)]
    )

    # 降低 similarity_top_k 以减少检索时间
    retriever = index.as_retriever(
        similarity_top_k=5,  # 降低从 10 到 5
        filters=filters,
    )

    # Custom system prompt
    system_prompt = """
    你是一个乐于助人的助手。
使用以下内容作为你所学习的知识，放在<context></context> XML标签内。
<context>
{{#context#}}
</context>
回答用户时：
如果你不知道，就直说你不知道。如果你在不确定的时候不知道，就寻求澄清。
避免提及你是从上下文中获取的信息。
并根据用户问题的语言来回答。
    """

    # Custom QA template
    from llama_index.core import PromptTemplate
    qa_template = PromptTemplate(
        system_prompt + "\n\n" + "Context information is below.\n---------------------\n{context_str}\n---------------------\nGiven the context information and not prior knowledge, answer the following question.\nQuestion: {query_str}\nAnswer: "
    )

    # 创建 Query Engine（不带 rerank，减少延迟）
    query_engine = index.as_query_engine(
        text_qa_template=qa_template,
        filters=filters,
    )

    # 缓存 query engine
    _query_engine_cache[cache_key] = query_engine
    return query_engine


def get_vector_index(table_name="llama_knowledge"):
    """
    Initializes and returns the VectorStoreIndex using PostgreSQL pgvector.
    Uses separate vector database configuration from settings.VECTOR_DATABASES.
    使用缓存提高性能。
    """
    # 检查缓存
    if table_name in _vector_index_cache:
        return _vector_index_cache[table_name]

    # Get vector database config from Django settings
    vector_db = settings.VECTOR_DATABASES['default']
    db_name = vector_db['NAME']
    host = vector_db['HOST']
    password = vector_db['PASSWORD']
    port = vector_db['PORT']
    user = vector_db['USER']

    # Initialize PGVectorStore
    # text-embedding-v2 embedding dimension is 1536
    vector_store = PGVectorStore.from_params(
        database=db_name,
        host=host,
        password=password,
        port=port,
        user=user,
        table_name=table_name,
        embed_dim=1536
    )

    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    # Return (or create) the index
    index = VectorStoreIndex.from_vector_store(
        vector_store=vector_store,
        storage_context=storage_context
    )

    # 缓存结果
    _vector_index_cache[table_name] = index
    return index


def ingest_document(file_path: str, doc_metadata: dict):
    """
    Loads a document from the file path, injects custom metadata (tenant_id, doc_id),
    and inserts it into the Vector Store.
    """
    import os

    print(f"[Ingest] Starting to process file: {file_path}")

    # 获取文件扩展名
    ext = os.path.splitext(file_path)[1].lower()
    print(f"[Ingest] File extension: {ext}")

    documents = []

    try:
        if ext == '.pdf':
            # PDF 文件
            try:
                from llama_index.readers.file import PDFReader
                reader = PDFReader()
                documents = reader.load_data(file=file_path)
                print(f"[Ingest] PDF loaded successfully, {len(documents)} pages")
            except Exception as e:
                print(f"[Ingest] PDFReader failed: {e}, trying pypdf")
                import pypdf
                from llama_index.core import Document
                reader = pypdf.PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages])
                documents = [Document(text=text)]
                print(f"[Ingest] pypdf loaded, text length: {len(text)}")

        elif ext in ['.docx', '.doc']:
            # Word 文件
            try:
                from llama_index.readers.file import DocxReader
                reader = DocxReader()
                documents = reader.load_data(file_path)
                print(f"[Ingest] Docx loaded successfully")
            except Exception as e:
                print(f"[Ingest] DocxReader failed: {e}, trying python-docx")
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                from llama_index.core import Document
                documents = [Document(text=text)]
                print(f"[Ingest] python-docx loaded, text length: {len(text)}")

        elif ext in ['.txt', '.md', '.json']:
            # 文本文件
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            from llama_index.core import Document
            documents = [Document(text=text)]
            print(f"[Ingest] Text file loaded, text length: {len(text)}")

        else:
            # 默认使用 SimpleDirectoryReader
            documents = SimpleDirectoryReader(input_files=[file_path]).load_data()
            print(f"[Ingest] SimpleDirectoryReader loaded {len(documents)} documents")

    except Exception as e:
        print(f"[Ingest] Error reading file: {e}")
        import traceback
        traceback.print_exc()
        raise

    if not documents:
        print(f"[Ingest] Warning: No documents extracted!")
        return False

    # 2. Inject enterprise metadata (for data isolation and filtering)
    for doc in documents:
        doc.metadata.update(doc_metadata)

    # 3. Write into Vector Store (pgVector)
    print(f"[Ingest] Inserting {len(documents)} documents into vector store")
    index = get_vector_index()
    for doc in documents:
        index.insert(doc)

    print(f"[Ingest] Completed successfully!")
    return True


def ask_question(query: str, tenant_id: str, chat_history=None, use_rerank=True):
    """
    优化后的查询函数 - 使用缓存的 Query Engine

    参数:
        query: 用户问题
        tenant_id: 租户ID（用于数据隔离）
        chat_history: 聊天历史（暂未使用）
        use_rerank: 是否使用 rerank（关闭可提升速度）

    性能优化:
    - 使用缓存的 Query Engine，避免重复创建
    - similarity_top_k 从 10 降低到 5
    - 可选关闭 rerank 以提升响应速度
    """
    # 使用缓存的 Query Engine
    query_engine = get_query_engine(tenant_id)

    # 直接查询
    response = query_engine.query(query)

    rerank_top_n = 5

    # 可选的 rerank
    if use_rerank and HAS_DASHSCOPE_RERANK and response.source_nodes and len(response.source_nodes) > 1:
        try:
            # 准备 rerank 文档
            documents = [node.node.get_content() for node in response.source_nodes]
            api_key = os.getenv('DASHSCOPE_API_KEY')
            if api_key:
                dashscope.api_key = api_key
                # 调用 DashScope rerank API
                rerank_result = TextReRank.call(
                    model='gte-rerank-v2',
                    query=query,
                    documents=documents,
                    top_n=rerank_top_n
                )
                if rerank_result.status_code == 200:
                    print(f"[Rerank] Success - got {len(rerank_result.output['results'])} results")
                    # 根据 rerank 结果重新排序 source_nodes
                    reranked_indices = rerank_result.output['results']
                    reranked_nodes = []
                    seen = set()
                    for r in reranked_indices:
                        idx = r['index']
                        if idx < len(response.source_nodes) and idx not in seen:
                            node = response.source_nodes[idx]
                            # relevance_score 已经是 0-1 的分数，越高越相关，直接使用
                            node.score = r['relevance_score']
                            reranked_nodes.append(node)
                            seen.add(idx)
                    # 添加未在 rerank 结果中的节点
                    for i, node in enumerate(response.source_nodes):
                        if i not in seen:
                            reranked_nodes.append(node)
                    response.source_nodes = reranked_nodes
        except Exception as e:
            print(f"Custom rerank failed: {e}")

    # Build sources JSON summary
    source_info = []
    response_str = str(response).strip()

    if response.source_nodes:
        from rag_app.models import Document
        for node in response.source_nodes:
            content = node.node.get_content()[:200]
            # Fix encoding issues
            try:
                content = content.encode('utf-8', errors='replace').decode('utf-8')
            except:
                content = str(content)

            # Get doc_id from metadata and resolve actual filename from database
            doc_id = node.node.metadata.get('doc_id', '')
            original_filename = ''
            if doc_id:
                try:
                    doc = Document.objects.get(id=doc_id)
                    original_filename = doc.filename
                except Document.DoesNotExist:
                    pass

            metadata = dict(node.node.metadata)
            if original_filename:
                metadata['original_filename'] = original_filename

            source_info.append({
                'score': node.score,
                'metadata': metadata,
                'content_preview': content
            })

    # Fix response encoding
    answer = str(response)
    try:
        answer = answer.encode('utf-8', errors='replace').decode('utf-8')
    except:
        pass

    # If still empty after all attempts, return a helpful message
    if not answer or answer.strip() == "Empty Response":
        if not source_info:
            answer = "抱歉，您的知识库中没有任何文档。请先在知识库中上传文档，然后我可以回答您关于这些文档的问题。"
        else:
            answer = "抱歉，我无法从您提供的文档中找到相关的答案。"

    return answer, source_info


def get_relevant_content(query: str, tenant_id: str, top_k: int = 5):
    """
    获取与查询相关的文档内容片段

    Args:
        query: 查询内容
        tenant_id: 租户ID
        top_k: 返回结果数量

    Returns:
        相关文档内容列表
    """
    from llama_index.core.vector_stores import MetadataFilters, ExactMatchFilter

    index = get_vector_index()

    # Data Isolation
    filters = MetadataFilters(
        filters=[ExactMatchFilter(key="tenant_id", value=tenant_id)]
    )

    # Retriever
    retriever = index.as_retriever(
        similarity_top_k=top_k,
        filters=filters,
    )

    # Get relevant nodes
    nodes = retriever.retrieve(query)

    results = []
    for i, node in enumerate(nodes):
        content = node.node.get_content()
        try:
            content = content.encode('utf-8', errors='replace').decode('utf-8')
        except:
            content = str(content)

        results.append({
            'rank': i + 1,
            'score': node.score,
            'doc_id': node.node.metadata.get('doc_id', ''),
            'file_name': node.node.metadata.get('file_name', ''),
            'page_label': node.node.metadata.get('page_label', ''),
            'content': content
        })

    return results


def get_document_content(tenant_id: str, doc_id: str):
    """
    获取指定文档的所有内容片段

    Args:
        tenant_id: 租户ID
        doc_id: 文档ID

    Returns:
        文档内容列表
    """
    from llama_index.core.vector_stores import MetadataFilters, ExactMatchFilter

    index = get_vector_index()

    # Filter by tenant_id and doc_id
    filters = MetadataFilters(
        filters=[
            ExactMatchFilter(key="tenant_id", value=tenant_id),
            ExactMatchFilter(key="doc_id", value=doc_id)
        ]
    )

    # Retriever with high top_k to get all chunks
    retriever = index.as_retriever(
        similarity_top_k=100,
        filters=filters,
    )

    # Use doc_id as query to retrieve all chunks of this document
    # Or use the filename pattern
    query = f"doc_id:{doc_id}"
    nodes = retriever.retrieve(query)

    # If no results, try using the doc_id directly
    if not nodes:
        nodes = retriever.retrieve(doc_id)

    results = []
    for i, node in enumerate(nodes):
        content = node.node.get_content()
        try:
            content = content.encode('utf-8', errors='replace').decode('utf-8')
        except:
            content = str(content)

        results.append({
            'chunk_id': i + 1,
            'score': node.score,
            'page_label': node.node.metadata.get('page_label', ''),
            'content': content
        })

    return results


def delete_document_from_vector(doc_id: str, tenant_id: str):
    """
    从向量数据库中删除指定文档的所有内容
    直接删除 data_llama_knowledge 表中的数据
    """
    import psycopg

    # Get vector database config
    vector_db = settings.VECTOR_DATABASES['default']
    conn = psycopg.connect(
        host=vector_db['HOST'],
        port=vector_db['PORT'],
        user=vector_db['USER'],
        password=vector_db['PASSWORD'],
        dbname=vector_db['NAME']
    )

    try:
        with conn.cursor() as cur:
            # 删除匹配 doc_id 的记录 (在 metadata 任意位置)
            cur.execute(
                """DELETE FROM data_llama_knowledge
                   WHERE metadata_->>'tenant_id' = %s
                   AND metadata_::text LIKE %s""",
                (tenant_id, f'%doc_id%{doc_id}%')
            )
            deleted_count = cur.rowcount

            conn.commit()
            return deleted_count
    finally:
        conn.close()
